import json

import docx
from docx.shared import Inches, RGBColor, Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

import csv

file_name = ""

# night order on the first night
first_night_order = []
# night order on other nights
other_night_order = []

# Character names that appear on night one
night_one_characters = []
# Character names that appear on other nights
night_x_characters = []
# Global file with the json data from night.json
night_file = json.load(open('night.json', 'r'))
night_file_name = "night.json"


# This will populate the night_one_characters and night_x_characters arrays
def init_characters():
    for i in night_file["first_night_characters"]:
        night_one_characters.append(i)

    for i in night_file["other_night_characters"]:
        night_x_characters.append(i)


# This will scan through the passed in .json file and if there is a first night
# wake, add that to the first night order. If there is a wake on other nights,
# add that to the other night order
def found_character(character_name):
    if character_name in night_one_characters:
        first_night_order.append(night_file["first_night"][character_name])

    if character_name in night_x_characters:
        other_night_order.append(night_file["other_night"][character_name])


# This will scan a script's json file provided in the same format as the
# official script tool return a list of all the characters in the script
def search_script(script_file_name):
    script_file = json.load(open(script_file_name))
    characters_appearing = []

    first = True

    for i in script_file:
        if not first:
            characters_appearing.append(i)
        else:
            first = False

    return characters_appearing


def transform_character_name(character_name):
    character_name = character_name.lower()
    character_name = character_name.replace(" ", "")
    character_name = character_name.replace("'", "")
    return character_name


# This method sorts the night orders by order_pos, using Selection Sort
def sort_night_order():
    for i in range(len(first_night_order)):
        min = first_night_order[i][3]['order_pos']
        min_index = i
        for j in range(i + 1, len(first_night_order)):
            if first_night_order[j][3]['order_pos'] < min:
                min = first_night_order[j][3]['order_pos']
                min_index = j

        first_night_order[i], first_night_order[min_index] = first_night_order[
                                                                 min_index], first_night_order[i]

    for i in range(len(other_night_order)):
        min = other_night_order[i][3]['order_pos']
        min_index = i
        for j in range(i + 1, len(other_night_order)):
            if other_night_order[j][3]['order_pos'] < min:
                min = other_night_order[j][3]['order_pos']
                min_index = j

        other_night_order[i], other_night_order[min_index] = other_night_order[
                                                                 min_index], other_night_order[i]


# This will create the night order and output the night order
def create_night_order(file_name_begin):
    # adding the boiler plate night order spots in every script
    found_character("dusk")
    found_character("dawn")

    # adding full script info if it's a full script
    if (int(input("Is this a Teensyville (0) or a full script (1): "))) == 1:
        found_character("demon_info")
        found_character("minion_info")
        found_character("travelers")

    sort_night_order()

    with open("Output_CSV_files/" + file_name_begin + "_Night_Order.csv",
              'w', encoding="utf-8") as f:
        for i in first_night_order:
            f.write(i[0]["name"] + " ,|," + '"' + i[1]["description"] + '"\n')
        f.write("\n")

        for i in other_night_order:
            f.write(i[0]["name"] + " ,|," + '"' + i[1]["description"] + '"\n')


def create_chart(file_name):
    characters = search_script("Input_Json_files/" + file_name)
    file_name = file_name.replace(".json", "")
    print(file_name)
    for character in characters:
        found_character(character)

    create_night_order(file_name)


def add_character():
    name = input("New Character Name: ")
    nights = int(input("1) First Night, 2) Other Nights, 3) Both: "))
    first_night_description = ""
    first_night_order_pos = ""
    other_night_description = ""
    other_night_order_pos = ""

    if nights != 2:
        first_night_description = input("New Character First Night Description: ")
        first_night_order_pos = transform_character_name(
            input("What character do they go after in first night order: "))
    if nights != 1:
        other_night_description = input("New Character Other Night Description: ")
        other_night_order_pos = transform_character_name(
            input("What character do they go after in other night order: "))

    first_night_alignment = input("New Character alignment: ")

    if first_night_order_pos != "":
        spot_to_add_character = night_one_characters.index(first_night_order_pos)
        night_one_characters.insert(spot_to_add_character + 1,
                                    transform_character_name(name))
        first_night_addition = [{
            "name": name.title()
        }, {
            "description": first_night_description
        }, {
            'default_alignment': first_night_alignment.title()
        }, {
            "order_pos":
                night_one_characters.index(transform_character_name(name)) + 1
        }]

        for character in night_file["first_night"]:
            night_file["first_night"][character][3][
                "order_pos"] = night_one_characters.index(character) + 1

        night_file["first_night"].update(
            {transform_character_name(name): first_night_addition})
        night_file["first_night_characters"].insert(spot_to_add_character + 1,
                                                    transform_character_name(name))

    if other_night_order_pos != "":
        spot_to_add_character = night_x_characters.index(other_night_order_pos)
        night_x_characters.insert(spot_to_add_character + 1,
                                  transform_character_name(name))
        other_night_addition = [{
            "name": name.title()
        }, {
            "description": other_night_description
        }, {
            'default_alignment': first_night_alignment.title()
        }, {
            "order_pos":
                night_x_characters.index(transform_character_name(name)) + 1
        }]

        for character in night_file["other_night"]:
            night_file["other_night"][character][3][
                "order_pos"] = night_x_characters.index(character) + 1

        night_file["other_night"].update(
            {transform_character_name(name): other_night_addition})
        night_file["other_night_characters"].insert(spot_to_add_character + 1,
                                                    transform_character_name(name))

    with open(night_file_name, 'w') as f:
        json.dump(night_file, f, indent=4)


def format_text(text, columnNumber):
    if columnNumber != 2:
        text.bold = True

    if columnNumber == 0:
        for i in first_night_order:
            if i[0]["name"] in text.text:
                if i[2]["default_alignment"] == "Good":
                    text.font.color.rgb = RGBColor(0x06, 0x63, 0xB9)
                elif i[2]["default_alignment"] == "Evil":
                    text.font.color.rgb = RGBColor(0xBF, 0x00, 0x00)

        for i in other_night_order:
            if i[0]["name"] in text.text:
                if i[2]["default_alignment"] == "Good":
                    text.font.color.rgb = RGBColor(0x06, 0x63, 0xB9)
                elif i[2]["default_alignment"] == "Evil":
                    text.font.color.rgb = RGBColor(0xBF, 0x00, 0x00)


def createDocument(filename):
    doc = docx.Document()

    # Each section needs to be formatted to have the right margins and height
    # Half an inch on each side and 14 inch height (Legal)
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.page_height = Inches(14)

    # The full 14 inches isn't printable, and this also allows users to say they only want the
    # top half of their paper printed from the chart
    printable_height = float(input("List in inches how tall your printable page is: "))
    with open("Output_CSV_files/" + filename + "_Night_Order.csv") as f:
        csv_reader = csv.reader(f)
        first_night_instructions = []
        other_night_instructions = []
        first_night_rows = True

        # This reads in the already parsed instructions from the csv file and saves them to
        # the appropriate file.
        for row in csv_reader:
            if row == []:
                first_night_rows = False
                continue
            if first_night_rows:
                first_night_instructions.append(row)
            else:
                other_night_instructions.append(row)

    night_one_header = doc.add_paragraph()
    night_one_header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    n1runner = night_one_header.add_run(filename + " - First Night")
    n1runner.bold = True
    n1runner.font.size = Pt(16)

    # Initializing a table with 3 columns and column widths that make sense for the content going in each column.
    table_first_night = doc.add_table(rows=0, cols=3)
    table_first_night.allow_autofit = False
    table_first_night.columns[0].width = Inches(1.2)
    table_first_night.columns[1].width = Inches(0.2)
    table_first_night.columns[2].width = Inches(6.1)

    shaded = True
    # Add each set of instructions to the table and format it correctly
    for instruction in first_night_instructions:
        row_cells = table_first_night.add_row().cells
        for i in range(3):
            text = row_cells[i].paragraphs[0].add_run(instruction[i])
            format_text(text, i)
            row_cells[i].paragraphs[0].style.paragraph_format.space_after = 0
            row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if i == 0:
                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if shaded:
                shading = parse_xml(r'<w:shd {} w:fill="D3D3D3"/>'.format(nsdecls('w')))
                row_cells[i]._tc.get_or_add_tcPr().append(shading)

        shaded = not shaded

    # Set the row height dependent on the printable height over the number of instructions
    for row in table_first_night.rows:
        row.height = Inches(printable_height / len(first_night_instructions))

    # No matter what the other night is printed on a different page than the first night
    doc.add_page_break()
    night_other_header = doc.add_paragraph()
    night_other_header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    onrunner = night_other_header.add_run(filename + " - Other Nights")
    onrunner.bold = True
    onrunner.font.size = Pt(16)

    # Initializing a table with 3 columns and column widths that make sense for the content going in each column.
    table_other_night = doc.add_table(rows=0, cols=3)
    table_other_night.allow_autofit = False
    table_other_night.columns[0].width = Inches(1.2)
    table_other_night.columns[1].width = Inches(0.2)
    table_other_night.columns[2].width = Inches(6.1)

    shaded = True
    # Add each set of instructions to the table and format it correctly
    for instruction in other_night_instructions:
        row_cells = table_other_night.add_row().cells
        for i in range(3):
            text = row_cells[i].paragraphs[0].add_run(instruction[i])
            format_text(text, i)
            row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if i == 0:
                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if shaded:
                shading = parse_xml(r'<w:shd {} w:fill="D3D3D3"/>'.format(nsdecls('w')))
                row_cells[i]._tc.get_or_add_tcPr().append(shading)
        shaded = not shaded

    # Set the row height dependent on the printable height over the number of instructions
    for row in table_other_night.rows:
        row.height = Inches(printable_height / len(other_night_instructions))

    # Save the document to Output_Night_Order_Sheets/<SCRIPT>_Formatted_Night.docx
    doc.save("Output_Night_Order_Sheets/" + filename + "_Formatted_Night.docx")


choice = int(input("Would you like to 1) add a character to the night order or 2) create a night order chart: "))

if int(input("Do you want to do this to a custom character list (1) or the default (0): ")) != 0:
    night_file_name = input("Enter the custom character list filename: ")
    night_file = json.load(open(night_file_name, 'r'))


init_characters()

if choice == 1:
    add_character()
else:
    file_name = input("Enter the script file name: ")
    create_chart(file_name)
    file_name = file_name.replace(".json", "")
    createDocument(file_name)
